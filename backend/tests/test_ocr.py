"""Integration tests for OCR/Text Workstation v1 routing and persistence."""

from __future__ import annotations

import hashlib
import json
from collections.abc import Iterator
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from PIL import Image
from pypdf import PdfReader, PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from app.api.ocr import provide_ocr_adapter
from app.config import Settings
from app.main import app
from app.services.artifact_service import get_latest_quality_report_artifact
from app.services.layout_text import build_pdf_layout_blocks
from app.services.ocr_adapters import (
    OcrExtractionResult,
    OcrLayoutLine,
    OcrLineMetric,
    OcrUnavailableError,
)
from app.services.pdf_renderer import get_pdf_renderer


class FakeOcrAdapter:
    def __init__(self) -> None:
        self.outputs: list[str] = []
        self.confidences: list[float | None] = []
        self.line_confidences: list[tuple[OcrLineMetric, ...]] = []
        self.layout_lines: list[tuple[OcrLayoutLine, ...]] = []
        self.image_sizes: list[tuple[int, int] | None] = []
        self.calls: list[Path] = []
        self.unavailable = False

    def extract_text(self, image_path: Path) -> str:
        return self.extract_result(image_path).text

    def extract_result(self, image_path: Path) -> OcrExtractionResult:
        self.calls.append(image_path)
        if self.unavailable:
            raise OcrUnavailableError
        index = len(self.calls) - 1
        return OcrExtractionResult(
            text=self.outputs[index],
            confidence=self.confidences[index] if index < len(self.confidences) else None,
            line_confidences=(
                self.line_confidences[index] if index < len(self.line_confidences) else ()
            ),
            layout_lines=(self.layout_lines[index] if index < len(self.layout_lines) else ()),
            image_width=(
                self.image_sizes[index][0]
                if index < len(self.image_sizes) and self.image_sizes[index] is not None
                else None
            ),
            image_height=(
                self.image_sizes[index][1]
                if index < len(self.image_sizes) and self.image_sizes[index] is not None
                else None
            ),
        )

    def tool_versions(self) -> dict[str, str]:
        return {"paddleocr": "test"}


class FakePdfRenderer:
    def __init__(self) -> None:
        self.calls: list[int] = []
        self.output_directories: list[Path] = []
        self.fail = False

    def render_page(self, pdf_path: Path, page_number: int, output_dir: Path) -> Path:
        self.calls.append(page_number)
        self.output_directories.append(output_dir)
        if self.fail:
            raise RuntimeError("simulated rendering failure")
        rendered = output_dir / f"page-{page_number}.png"
        rendered.write_bytes(b"fake rendered page")
        return rendered


@pytest.fixture(autouse=True)
def _allow_larger_ocr_fixtures(settings: Settings) -> None:
    settings.max_upload_bytes = 2 * 1024 * 1024


@pytest.fixture
def ocr_fakes(client: TestClient) -> Iterator[tuple[FakeOcrAdapter, FakePdfRenderer]]:
    adapter = FakeOcrAdapter()
    renderer = FakePdfRenderer()
    app.dependency_overrides[provide_ocr_adapter] = lambda: adapter
    app.dependency_overrides[get_pdf_renderer] = lambda: renderer
    yield adapter, renderer


def _upload(client: TestClient, name: str, content: bytes, content_type: str) -> dict[str, object]:
    response = client.post("/api/uploads", files={"file": (name, content, content_type)})
    assert response.status_code == 201
    return response.json()


def _upload_and_audit(
    client: TestClient, name: str, content: bytes, content_type: str
) -> tuple[dict[str, object], dict[str, object]]:
    upload = _upload(client, name, content, content_type)
    response = client.post(f"/api/documents/{upload['id']}/audit")
    assert response.status_code == 201
    return upload, response.json()


def _pdf_pages_bytes(*page_texts: str | None) -> bytes:
    writer = PdfWriter()
    for text in page_texts:
        page = writer.add_blank_page(width=200, height=200)
        if text is not None:
            font = DictionaryObject(
                {
                    NameObject("/Type"): NameObject("/Font"),
                    NameObject("/Subtype"): NameObject("/Type1"),
                    NameObject("/BaseFont"): NameObject("/Helvetica"),
                }
            )
            page[NameObject("/Resources")] = DictionaryObject(
                {
                    NameObject("/Font"): DictionaryObject(
                        {NameObject("/F1"): writer._add_object(font)}
                    )
                }
            )
            stream = DecodedStreamObject()
            stream.set_data(f"BT /F1 12 Tf 10 100 Td ({text}) Tj ET".encode())
            page[NameObject("/Contents")] = writer._add_object(stream)
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def _docx_bytes(*paragraphs: str) -> bytes:
    document = DocxDocument()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _docx_with_table_bytes() -> bytes:
    """A paragraph, a 2x2 table, then a paragraph — to exercise ordered table extraction."""
    document = DocxDocument()
    document.add_paragraph("Intro")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "R1C1"
    table.rows[0].cells[1].text = "R1C2"
    table.rows[1].cells[0].text = "R2C1"
    table.rows[1].cells[1].text = "R2C2"
    document.add_paragraph("Outro")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _image_bytes(image_format: str, size: tuple[int, int] = (10, 10)) -> bytes:
    buffer = BytesIO()
    Image.new("RGB", size, color="white").save(buffer, format=image_format)
    return buffer.getvalue()


def _artifact_path(
    document_data_dir: Path, document_id: object, artifact_id: object
) -> Path:
    return document_data_dir / str(document_id) / "artifacts" / f"{artifact_id}.json"


def _artifact_payloads_by_type(
    document_data_dir: Path, document_id: object, artifact_type: str
) -> list[dict[str, object]]:
    directory = document_data_dir / str(document_id) / "artifacts"
    payloads = [
        json.loads(path.read_text(encoding="utf-8"))
        for path in directory.glob("*.json")
    ]
    return [payload for payload in payloads if payload.get("artifact_type") == artifact_type]


def _set_audit_page_fields(
    document_data_dir: Path,
    document_id: object,
    audit_id: object,
    page_index: int,
    **fields: object,
) -> None:
    """Overwrite persisted audit page fields to stage a specific per-page routing decision."""
    path = _artifact_path(document_data_dir, document_id, audit_id)
    payload = json.loads(path.read_text(encoding="utf-8"))
    payload["content"]["pages"][page_index].update(fields)
    path.write_text(json.dumps(payload), encoding="utf-8")


_GOOD_TEXT = "The quick brown fox jumps over the lazy dog near the calm winding river today"


def test_pdf_text_layer_creates_text_artifact_without_ocr(
    client: TestClient,
    settings: Settings,
    document_data_dir: Path,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, renderer = ocr_fakes
    upload, audit = _upload_and_audit(
        client, "text.pdf", _pdf_pages_bytes("Digital text"), "application/pdf"
    )
    audit_path = _artifact_path(document_data_dir, upload["id"], audit["id"])
    audit_bytes = audit_path.read_bytes()

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    artifact = response.json()
    content = artifact["content"]
    assert artifact["artifact_type"] == "text_result"
    assert artifact["station"] == "ocr"
    assert artifact["input_artifact_id"] == upload["original_artifact"]["id"]
    assert artifact["input_audit_artifact_id"] == audit["id"]
    assert content["source"] == "pdf_text_layer"
    assert content["text"] == "Digital text"
    assert content["text_char_count"] == len(content["text"])
    assert content["readable_text"] == "Digital text"
    assert content["reading_text"] == "Digital text"
    assert content["reading_text_version"] == "1"
    assert content["reading_text_status"] == "heuristic"
    assert content["pages"] == [
        {
            "page_number": 1,
            "source": "pdf_text_layer",
            "has_text_layer": True,
            "ocr_used": False,
            "text": "Digital text",
            "text_char_count": len("Digital text"),
            "ocr_confidence": None,
            "ocr_line_confidences": [],
        }
    ]
    assert content["tool_versions"]["pypdf"]
    assert adapter.calls == []
    assert renderer.calls == []
    assert _artifact_path(document_data_dir, upload["id"], artifact["id"]).is_file()
    quality_report = get_latest_quality_report_artifact(settings, str(upload["id"]))
    assert quality_report is not None
    assert quality_report.artifact_type == "quality_report"
    assert quality_report.station == "ocr_quality"
    assert quality_report.media_type == "application/json"
    assert quality_report.input_artifact_id == artifact["input_artifact_id"]
    assert quality_report.input_audit_artifact_id == audit["id"]
    assert quality_report.input_text_artifact_id == artifact["id"]
    summary = quality_report.content
    assert summary.quality_report_version == "1"
    assert summary.page_count == 1
    assert summary.text_layer_pages == 1
    assert summary.ocr_pages == 0
    assert summary.mixed_source is False
    assert summary.text_source == "pdf_text_layer"
    assert summary.pages_needing_ocr == 0
    assert summary.ocr_pages_with_confidence == 0
    assert summary.ocr_lines_with_confidence == 0
    assert summary.ocr_page_confidence_mean is None
    assert summary.final_char_count == len("Digital text")
    assert summary.final_word_count == 2
    assert summary.pages_without_text == 0
    assert "text" not in summary.model_dump()
    assert "Digital text" not in quality_report.model_dump_json()
    assert audit_path.read_bytes() == audit_bytes


def test_mixed_pdf_routes_each_page_and_preserves_order(
    client: TestClient,
    upload_dir: Path,
    document_data_dir: Path,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, renderer = ocr_fakes
    adapter.outputs = ["Scan two", "Scan three"]
    upload, _ = _upload_and_audit(
        client,
        "mixed.pdf",
        _pdf_pages_bytes("Digital one", None, None),
        "application/pdf",
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    content = response.json()["content"]
    assert content["source"] == "pdf_mixed"
    assert content["text"] == "Digital one\n\nScan two\n\nScan three"
    assert content["readable_text"] == (
        "Digital one\n\n----- page 2 -----\n\nScan two\n\n----- page 3 -----\n\nScan three"
    )
    assert [page["source"] for page in content["pages"]] == [
        "pdf_text_layer",
        "paddleocr",
        "paddleocr",
    ]
    assert renderer.calls == [2, 3]
    assert [path.name for path in adapter.calls] == ["page-2.png", "page-3.png"]
    assert all(path.parent == Path("/tmp") for path in renderer.output_directories)
    assert all(upload_dir not in path.parents for path in adapter.calls)
    assert all(not path.exists() for path in adapter.calls)
    artifact_directory = document_data_dir / str(upload["id"]) / "artifacts"
    assert all(path.suffix == ".json" for path in artifact_directory.rglob("*"))


def test_docx_extracts_paragraphs_without_ocr(
    client: TestClient,
    settings: Settings,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, renderer = ocr_fakes
    upload, _ = _upload_and_audit(
        client,
        "document.docx",
        _docx_bytes("First", "Second"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    content = response.json()["content"]
    assert content["source"] == "docx_text"
    assert content["text"] == "First\nSecond"
    assert content["readable_text"] == "First Second"
    assert content["pages"] == []
    assert adapter.calls == []
    assert renderer.calls == []
    quality_report = get_latest_quality_report_artifact(settings, str(upload["id"]))
    assert quality_report is not None
    assert quality_report.content.text_source == "docx_text"
    assert quality_report.content.page_count == 0
    assert quality_report.content.text_layer_pages == 0
    assert quality_report.content.ocr_pages == 0
    assert quality_report.content.final_word_count == 2


def test_docx_extracts_table_cells_in_document_order(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, renderer = ocr_fakes
    upload, _ = _upload_and_audit(
        client,
        "tables.docx",
        _docx_with_table_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    content = response.json()["content"]
    assert content["source"] == "docx_text"
    assert content["structured_content_version"] == "1"
    table = content["structured_content"]["pages"][0]["tables"][0]
    assert (table["row_count"], table["column_count"]) == (2, 2)
    assert all("text" not in cell for cell in table["cells"])
    # Deterministic order: leading paragraph, table rows (cells tab-joined, rows newline-joined),
    # trailing paragraph. Table cell text must be present — paragraph-only extraction dropped it.
    assert content["text"] == "Intro\nR1C1\tR1C2\nR2C1\tR2C2\nOutro"
    assert content["readable_text"] == "Intro R1C1 R1C2 R2C1 R2C2 Outro"
    assert content["text_char_count"] == len(content["text"])
    assert content["pages"] == []
    assert adapter.calls == []
    assert renderer.calls == []


@pytest.mark.parametrize(
    ("extension", "mime_type", "image_format"),
    [("png", "image/png", "PNG"), ("jpg", "image/jpeg", "JPEG")],
)
def test_image_uses_ocr_once(
    client: TestClient,
    settings: Settings,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
    extension: str,
    mime_type: str,
    image_format: str,
) -> None:
    adapter, renderer = ocr_fakes
    adapter.outputs = ["Image text"]
    adapter.confidences = [0.91]
    adapter.line_confidences = [(OcrLineMetric(1, 0.91, len("Image text")),)]
    upload, _ = _upload_and_audit(
        client, f"image.{extension}", _image_bytes(image_format), mime_type
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    content = response.json()["content"]
    assert content["source"] == "paddleocr"
    assert content["text"] == "Image text"
    assert content["readable_text"] == "Image text"
    assert content["pages"][0]["page_number"] == 1
    assert content["pages"][0]["ocr_confidence"] == 0.91
    assert content["pages"][0]["ocr_line_confidences"] == [
        {"line_index": 1, "confidence": 0.91, "text_char_count": len("Image text")}
    ]
    assert len(adapter.calls) == 1
    assert renderer.calls == []
    quality_report = get_latest_quality_report_artifact(settings, str(upload["id"]))
    assert quality_report is not None
    assert quality_report.content.page_count == 1
    assert quality_report.content.text_layer_pages == 0
    assert quality_report.content.ocr_pages == 1
    assert quality_report.content.ocr_pages_with_confidence == 1
    assert quality_report.content.ocr_lines_with_confidence == 1
    assert quality_report.content.ocr_page_confidence_mean == 0.91
    assert quality_report.content.flags == ["ocr_used"]
    assert "Image text" not in quality_report.model_dump_json()


def test_image_emits_structured_field_when_geometry_is_unavailable(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, _ = ocr_fakes
    adapter.outputs = ["Name: Max Mustermann"]
    upload, _ = _upload_and_audit(client, "form.png", _image_bytes("PNG"), "image/png")

    content = client.post(f"/api/documents/{upload['id']}/ocr").json()["content"]

    assert content["text"] == "Name: Max Mustermann"
    assert content["text_geometry"] is None
    field = content["structured_content"]["pages"][0]["fields"][0]
    assert field["label"] == "Name"
    assert field["field_type_hint"] == "person_name"
    assert field["source"] == "canonical_text"


def test_good_text_pdf_does_not_initialize_ocr(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, renderer = ocr_fakes
    upload, _ = _upload_and_audit(
        client, "good.pdf", _pdf_pages_bytes(_GOOD_TEXT), "application/pdf"
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    content = response.json()["content"]
    assert content["source"] == "pdf_text_layer"
    # A clean text-layer PDF never renders a page or touches the OCR adapter (no PaddleOCR init).
    assert adapter.calls == []
    assert renderer.calls == []


def test_broken_text_layer_page_routes_to_ocr(
    client: TestClient,
    document_data_dir: Path,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, renderer = ocr_fakes
    adapter.outputs = ["OCR recovered text"]
    upload, audit = _upload_and_audit(
        client, "broken.pdf", _pdf_pages_bytes("Digital text"), "application/pdf"
    )
    _set_audit_page_fields(
        document_data_dir,
        upload["id"],
        audit["id"],
        0,
        text_quality_status="BROKEN_TEXT_LAYER",
        needs_ocr=True,
        recommended_text_source="ocr",
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    content = response.json()["content"]
    assert content["source"] == "paddleocr"
    page = content["pages"][0]
    assert page["source"] == "paddleocr"
    assert page["ocr_used"] is True
    assert page["has_text_layer"] is False
    assert page["text"] == "OCR recovered text"
    assert renderer.calls == [1]
    assert len(adapter.calls) == 1


def test_empty_text_layer_page_routes_to_ocr(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, renderer = ocr_fakes
    adapter.outputs = ["Scanned page text"]
    adapter.confidences = [0.84]
    adapter.line_confidences = [(OcrLineMetric(1, 0.84, len("Scanned page text")),)]
    upload, _ = _upload_and_audit(
        client, "scan.pdf", _pdf_pages_bytes(None), "application/pdf"
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    content = response.json()["content"]
    assert content["source"] == "paddleocr"
    assert content["pages"][0]["ocr_used"] is True
    assert content["pages"][0]["ocr_confidence"] == 0.84
    assert content["pages"][0]["ocr_line_confidences"][0]["line_index"] == 1
    assert "text" not in content["pages"][0]["ocr_line_confidences"][0]
    assert content["text"] == "Scanned page text"
    assert content["pages"][0]["text"] == "Scanned page text"
    assert renderer.calls == [1]
    assert len(adapter.calls) == 1


def test_mixed_quality_pdf_routes_each_page(
    client: TestClient,
    settings: Settings,
    document_data_dir: Path,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, renderer = ocr_fakes
    adapter.outputs = ["Broken page OCR", "Scanned page OCR"]
    adapter.confidences = [0.8, 0.6]
    adapter.line_confidences = [
        (OcrLineMetric(1, 0.8, len("Broken page OCR")),),
        (OcrLineMetric(1, 0.6, len("Scanned page OCR")),),
    ]
    upload, audit = _upload_and_audit(
        client, "mixed.pdf", _pdf_pages_bytes(_GOOD_TEXT, _GOOD_TEXT, None), "application/pdf"
    )
    # Page 1 stays GOOD (text layer), page 2 is corrupted to BROKEN, page 3 is empty — the two
    # OCR-required pages must be rendered while the good page is not.
    _set_audit_page_fields(
        document_data_dir,
        upload["id"],
        audit["id"],
        1,
        text_quality_status="BROKEN_TEXT_LAYER",
        needs_ocr=True,
        recommended_text_source="ocr",
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    content = response.json()["content"]
    assert content["source"] == "pdf_mixed"
    assert [page["source"] for page in content["pages"]] == [
        "pdf_text_layer",
        "paddleocr",
        "paddleocr",
    ]
    assert renderer.calls == [2, 3]
    assert len(adapter.calls) == 2
    quality_report = get_latest_quality_report_artifact(settings, str(upload["id"]))
    assert quality_report is not None
    summary = quality_report.content
    assert summary.page_count == 3
    assert summary.text_layer_pages == 1
    assert summary.ocr_pages == 2
    assert summary.mixed_source is True
    assert summary.text_source == "pdf_mixed"
    assert summary.good_text_layer_pages == 1
    assert summary.broken_text_layer_pages == 1
    assert summary.empty_text_layer_pages == 1
    assert summary.pages_needing_ocr == 2
    assert summary.ocr_pages_with_confidence == 2
    assert summary.ocr_lines_with_confidence == 2
    assert summary.ocr_page_confidence_mean == pytest.approx(0.7)
    assert summary.ocr_page_confidence_min == 0.6
    assert summary.ocr_page_confidence_max == 0.8
    assert summary.flags == ["pdf_mixed", "ocr_used"]


def test_ocr_required_page_without_runtime_returns_503(
    client: TestClient,
    document_data_dir: Path,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, _ = ocr_fakes
    adapter.unavailable = True
    upload, audit = _upload_and_audit(
        client, "broken.pdf", _pdf_pages_bytes("Digital text"), "application/pdf"
    )
    _set_audit_page_fields(
        document_data_dir,
        upload["id"],
        audit["id"],
        0,
        text_quality_status="BROKEN_TEXT_LAYER",
        needs_ocr=True,
        recommended_text_source="ocr",
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 503
    # The broken text layer is never silently used as a result: no text_result artifact is written.
    artifact_directory = document_data_dir / str(upload["id"]) / "artifacts"
    assert [path.stem for path in artifact_directory.glob("*.json")] == [str(audit["id"])]


def test_legacy_audit_without_quality_fields_uses_text_layer(
    client: TestClient,
    document_data_dir: Path,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, renderer = ocr_fakes
    upload, audit = _upload_and_audit(
        client, "legacy.pdf", _pdf_pages_bytes("Digital text"), "application/pdf"
    )
    path = _artifact_path(document_data_dir, upload["id"], audit["id"])
    payload = json.loads(path.read_text(encoding="utf-8"))
    for key in (
        "text_quality_status",
        "text_quality_score",
        "text_quality_reasons",
        "recommended_text_source",
        "needs_ocr",
    ):
        payload["content"]["pages"][0].pop(key, None)
    path.write_text(json.dumps(payload), encoding="utf-8")

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    content = response.json()["content"]
    # Audits predating the quality gate carry no needs_ocr, so routing falls back to has_text_layer.
    assert content["source"] == "pdf_text_layer"
    assert adapter.calls == []
    assert renderer.calls == []


def test_missing_audit_returns_409(
    client: TestClient, ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer]
) -> None:
    upload = _upload(client, "blank.pdf", _pdf_pages_bytes(None), "application/pdf")

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 409


def test_legacy_metadata_without_original_artifact_returns_409(
    client: TestClient,
    document_data_dir: Path,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload = _upload(client, "text.pdf", _pdf_pages_bytes("Text"), "application/pdf")
    document_id = str(upload["id"])
    metadata_path = document_data_dir / document_id / "document.json"
    metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
    metadata.pop("original_artifact")
    metadata.pop("sha256")
    metadata.pop("detected_mime_type")
    metadata_path.write_text(json.dumps(metadata), encoding="utf-8")

    response = client.post(f"/api/documents/{document_id}/ocr")

    assert response.status_code == 409


def test_get_without_text_artifact_returns_404(client: TestClient) -> None:
    upload, _ = _upload_and_audit(
        client, "text.pdf", _pdf_pages_bytes("Text"), "application/pdf"
    )

    response = client.get(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 404


def test_get_returns_latest_text_artifact(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    upload, _ = _upload_and_audit(
        client, "text.pdf", _pdf_pages_bytes("Text"), "application/pdf"
    )
    timestamps = iter(["2026-07-01T10:00:00.000001Z", "2026-07-01T10:00:00.000002Z"])
    monkeypatch.setattr("app.services.ocr_service._now_utc_iso", lambda: next(timestamps))
    first = client.post(f"/api/documents/{upload['id']}/ocr")
    second = client.post(f"/api/documents/{upload['id']}/ocr")

    response = client.get(f"/api/documents/{upload['id']}/ocr")

    assert first.status_code == 201
    assert second.status_code == 201
    assert response.status_code == 200
    assert response.json()["id"] == second.json()["id"]


def test_rerun_creates_new_immutable_quality_report(
    client: TestClient,
    settings: Settings,
    document_data_dir: Path,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, _ = _upload_and_audit(
        client, "text.pdf", _pdf_pages_bytes("Stable text"), "application/pdf"
    )

    first_text = client.post(f"/api/documents/{upload['id']}/ocr")
    assert first_text.status_code == 201
    first_reports = _artifact_payloads_by_type(
        document_data_dir, upload["id"], "quality_report"
    )
    assert len(first_reports) == 1
    first_report = first_reports[0]
    first_report_path = _artifact_path(
        document_data_dir, upload["id"], first_report["id"]
    )
    first_report_bytes = first_report_path.read_bytes()
    first_text_path = _artifact_path(
        document_data_dir, upload["id"], first_text.json()["id"]
    )
    first_text_bytes = first_text_path.read_bytes()

    second_text = client.post(f"/api/documents/{upload['id']}/ocr")

    assert second_text.status_code == 201
    reports = _artifact_payloads_by_type(document_data_dir, upload["id"], "quality_report")
    assert len(reports) == 2
    assert first_report_path.read_bytes() == first_report_bytes
    assert first_text_path.read_bytes() == first_text_bytes
    assert {report["input_text_artifact_id"] for report in reports} == {
        first_text.json()["id"],
        second_text.json()["id"],
    }
    latest = get_latest_quality_report_artifact(settings, str(upload["id"]))
    assert latest is not None
    assert latest.input_text_artifact_id == second_text.json()["id"]


def test_hash_mismatch_prevents_text_artifact(
    client: TestClient,
    upload_dir: Path,
    document_data_dir: Path,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, audit = _upload_and_audit(
        client, "text.pdf", _pdf_pages_bytes("Text"), "application/pdf"
    )
    original = upload["original_artifact"]
    assert isinstance(original, dict)
    (upload_dir / str(original["storage_filename"])).write_bytes(b"tampered")

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 409
    artifact_directory = document_data_dir / str(upload["id"]) / "artifacts"
    assert [path.stem for path in artifact_directory.glob("*.json")] == [str(audit["id"])]


def test_audit_input_artifact_mismatch_returns_409(
    client: TestClient,
    document_data_dir: Path,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, audit = _upload_and_audit(
        client, "text.pdf", _pdf_pages_bytes("Text"), "application/pdf"
    )
    path = _artifact_path(document_data_dir, upload["id"], audit["id"])
    payload = json.loads(path.read_text(encoding="utf-8"))
    payload["input_artifact_id"] = "f" * 32
    payload["content"]["input_artifact_id"] = "f" * 32
    path.write_text(json.dumps(payload), encoding="utf-8")

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 409


def test_inconsistent_pdf_audit_page_list_returns_409(
    client: TestClient,
    document_data_dir: Path,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, audit = _upload_and_audit(
        client, "text.pdf", _pdf_pages_bytes("Text"), "application/pdf"
    )
    path = _artifact_path(document_data_dir, upload["id"], audit["id"])
    payload = json.loads(path.read_text(encoding="utf-8"))
    payload["content"]["pages"][0]["page_number"] = 2
    path.write_text(json.dumps(payload), encoding="utf-8")

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 409


def test_rendering_failure_returns_422(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, renderer = ocr_fakes
    renderer.fail = True
    upload, _ = _upload_and_audit(
        client, "blank.pdf", _pdf_pages_bytes(None), "application/pdf"
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 422
    assert adapter.calls == []


def test_paddleocr_unavailable_returns_503(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, _ = ocr_fakes
    adapter.unavailable = True
    upload, _ = _upload_and_audit(
        client, "image.png", _image_bytes("PNG"), "image/png"
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 503


def test_corrupt_original_returns_422_after_integrity_validation(
    client: TestClient,
    upload_dir: Path,
    document_data_dir: Path,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, _ = _upload_and_audit(
        client, "text.pdf", _pdf_pages_bytes("Text"), "application/pdf"
    )
    document_id = str(upload["id"])
    metadata_path = document_data_dir / document_id / "document.json"
    metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
    original_path = upload_dir / metadata["original_artifact"]["storage_filename"]
    corrupt = b"%PDF-1.4 broken after audit"
    digest = hashlib.sha256(corrupt).hexdigest()
    original_path.write_bytes(corrupt)
    metadata["sha256"] = digest
    metadata["original_artifact"]["sha256"] = digest
    metadata["size"] = len(corrupt)
    metadata["original_artifact"]["size_bytes"] = len(corrupt)
    metadata_path.write_text(json.dumps(metadata), encoding="utf-8")

    response = client.post(f"/api/documents/{document_id}/ocr")

    assert response.status_code == 422


def _pdf_two_column_table_bytes() -> bytes:
    """Synthetic single-page offer: two side-by-side blocks and a table. No real data.

    Text runs are placed at absolute page coordinates so pypdf's layout extraction can reconstruct
    columns and table rows; the default extraction linearises them.
    """
    runs = [
        (40, 792, 8, "Synthetic header"),
        (40, 770, 18, "KOSTENVORANSCHLAG"),
        (40, 730, 10, "AUFTRAGNEHMER"), (320, 730, 10, "AUFTRAGGEBER"),
        (40, 712, 10, "Sanierungsbau Perchtoldsdorf GmbH"),
        (320, 712, 10, "Herr Dipl.-Ing. Franz Hubermayr"),
        (40, 694, 10, "Lindenstrasse 42"), (320, 694, 10, "Rosengasse 7/12"),
        (40, 640, 10, "Angebot Nr.: KV-2026-0417"), (320, 640, 10, "Datum: 01.07.2026"),
        (40, 590, 10, "Pos."), (90, 590, 10, "Leistung"), (330, 590, 10, "Menge"),
        (390, 590, 10, "Einheit"), (450, 590, 10, "Einzelpreis"), (530, 590, 10, "Gesamt"),
        (40, 572, 10, "1"), (90, 572, 10, "Abbrucharbeiten Innenwaende"),
        (330, 572, 10, "45"), (390, 572, 10, "m2"), (450, 572, 10, "38,00"),
        (530, 572, 10, "1.710,00"), (40, 554, 10, "2"),
        (90, 554, 10, "Fassadendaemmung"), (330, 554, 10, "180"),
        (390, 554, 10, "m2"), (450, 554, 10, "92,00"),
        (530, 554, 10, "16.560,00"),
        (40, 100, 8, "Synthetic caption"),
        (40, 8, 8, "Synthetic footer"),
    ]
    writer = PdfWriter()
    page = writer.add_blank_page(width=600, height=800)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})}
    )
    data = "".join(
        f"BT /F1 {font_size} Tf {x} {y} Td ({text}) Tj ET\n"
        for x, y, font_size, text in runs
    )
    stream = DecodedStreamObject()
    stream.set_data(data.encode("latin-1"))
    page[NameObject("/Contents")] = writer._add_object(stream)
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def _line_with(text: str, *tokens: str) -> str | None:
    return next(
        (line for line in text.split("\n") if all(token in line for token in tokens)),
        None,
    )


def test_pdf_text_layer_produces_layout_text_result(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, renderer = ocr_fakes
    fixture = _pdf_two_column_table_bytes()
    fixture_page = PdfReader(BytesIO(fixture)).pages[0]
    first_blocks = build_pdf_layout_blocks(fixture_page, 1, fixture_page.extract_text() or "")
    second_blocks = build_pdf_layout_blocks(fixture_page, 1, fixture_page.extract_text() or "")
    assert first_blocks == second_blocks
    upload, _ = _upload_and_audit(client, "offer.pdf", fixture, "application/pdf")

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    content = response.json()["content"]
    # Technical raw text is the unchanged default extraction — the offset-stable PII input.
    expected_canonical = "\n\n".join(
        page.extract_text() or "" for page in PdfReader(BytesIO(fixture)).pages
    )
    assert content["text"] == expected_canonical
    layout = content["layout_text_result"]
    assert layout is not None
    # The layout rendering is a distinct, additive view — not the canonical text.
    assert layout != content["text"]
    # Two-column block stays side by side on one line.
    assert _line_with(layout, "AUFTRAGNEHMER", "AUFTRAGGEBER") is not None
    # Table header and its first value row each stay on a single readable line (no line-by-line
    # collapse of header vs values).
    assert _line_with(layout, "Pos.", "Leistung", "Gesamt") is not None
    assert _line_with(layout, "Abbrucharbeiten Innenwaende", "1.710,00") is not None
    blocks = content["layout_blocks"]
    assert content["layout_blocks_version"] == "1"
    assert [(block["page_number"], block["order"]) for block in blocks] == [
        (1, order) for order in range(1, len(blocks) + 1)
    ]
    assert all(0 <= block[key] <= 1 for block in blocks for key in ("x0", "y0", "x1", "y1"))
    assert any(block["block_type"] == "heading" for block in blocks)
    assert any(block["block_type"] == "header" for block in blocks)
    assert any(block["block_type"] == "caption" for block in blocks)
    assert any(block["block_type"] == "footer" for block in blocks)
    contractor = next(block for block in blocks if "Sanierungsbau" in block["text"])
    customer = next(block for block in blocks if "Franz Hubermayr" in block["text"])
    assert contractor["order"] < customer["order"]
    assert contractor["source"] == customer["source"] == "pdf_text_layer"
    structured = content["structured_content"]
    assert structured is not None
    # The canonical pypdf string safely exposes the first label/value pair. The horizontally
    # separated date is not split out unless its canonical boundary is unambiguous.
    assert "Angebot Nr." in {
        field["label"] for field in structured["pages"][0]["fields"]
    }
    # A clean text-layer PDF still never touches OCR.
    assert adapter.calls == []
    assert renderer.calls == []


def test_layout_text_result_absent_for_docx(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, _ = _upload_and_audit(
        client,
        "document.docx",
        _docx_bytes("First", "Second"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    content = response.json()["content"]
    assert content["layout_text_result"] is None
    assert content["layout_blocks_version"] == "1"
    assert content["layout_blocks"][0]["block_type"] == "fallback"


def test_layout_text_result_absent_for_image(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, _ = ocr_fakes
    adapter.outputs = ["Image text"]
    adapter.layout_lines = [
        (
            OcrLayoutLine(
                text="Image text",
                polygon=((10.0, 20.0), (190.0, 20.0), (190.0, 50.0), (10.0, 50.0)),
                confidence=0.9,
            ),
        )
    ]
    adapter.image_sizes = [(200, 100)]
    upload, _ = _upload_and_audit(client, "image.png", _image_bytes("PNG"), "image/png")

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    content = response.json()["content"]
    assert content["layout_text_result"] is None
    assert content["text"] == "Image text"
    assert content["pages"][0]["text"] == "Image text"
    assert content["layout_blocks"] == [
        {
            "page_number": 1,
            "order": 1,
            "block_type": "body",
            "text": "Image text",
            "x0": 0.05,
            "y0": 0.2,
            "x1": 0.95,
            "y1": 0.5,
            "source": "paddleocr",
            "confidence": 0.9,
        }
    ]


def test_layout_text_result_marks_ocr_pages_and_page_boundaries(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, _ = ocr_fakes
    adapter.outputs = ["Scanned page two text"]
    upload, _ = _upload_and_audit(
        client, "mixed.pdf", _pdf_pages_bytes(_GOOD_TEXT, None), "application/pdf"
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    layout = response.json()["content"]["layout_text_result"]
    assert layout is not None
    # Page 1 (text layer) contributes a layout rendering.
    assert _GOOD_TEXT.split()[0] in layout
    # Page boundary is visible and the OCR page is marked as not reconstructed, falling back to its
    # linear text.
    assert "----- page 2 -----" in layout
    assert "[page 2: layout not reconstructed]" in layout
    assert "Scanned page two text" in layout
    blocks = response.json()["content"]["layout_blocks"]
    assert {block["page_number"] for block in blocks} == {1, 2}
    page_two = [block for block in blocks if block["page_number"] == 2]
    assert page_two[0]["block_type"] == "fallback"
    assert page_two[0]["source"] == "fallback"


def test_legacy_text_artifact_without_additive_fields_remains_valid(
    client: TestClient,
    document_data_dir: Path,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, _ = _upload_and_audit(
        client, "text.pdf", _pdf_pages_bytes("Digital text"), "application/pdf"
    )
    created = client.post(f"/api/documents/{upload['id']}/ocr").json()
    path = _artifact_path(document_data_dir, upload["id"], created["id"])
    payload = json.loads(path.read_text(encoding="utf-8"))
    # Simulate a legacy artifact written before this field existed.
    payload["content"].pop("readable_text", None)
    payload["content"].pop("reading_text_version", None)
    payload["content"].pop("reading_text", None)
    payload["content"].pop("reading_text_status", None)
    payload["content"].pop("reading_text_flags", None)
    payload["content"].pop("reading_text_map_version", None)
    payload["content"].pop("reading_text_map", None)
    payload["content"].pop("layout_text_result", None)
    payload["content"].pop("layout_blocks_version", None)
    payload["content"].pop("layout_blocks", None)
    payload["content"].pop("structured_content_version", None)
    payload["content"].pop("structured_content", None)
    for page in payload["content"]["pages"]:
        page.pop("ocr_confidence", None)
        page.pop("ocr_line_confidences", None)
    path.write_text(json.dumps(payload), encoding="utf-8")

    response = client.get(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 200
    assert response.json()["content"]["readable_text"] is None
    assert response.json()["content"]["reading_text_version"] is None
    assert response.json()["content"]["reading_text"] is None
    assert response.json()["content"]["reading_text_status"] is None
    assert response.json()["content"]["reading_text_flags"] == []
    assert response.json()["content"]["layout_text_result"] is None
    assert response.json()["content"]["layout_blocks_version"] is None
    assert response.json()["content"]["layout_blocks"] == []
    assert response.json()["content"]["structured_content_version"] is None
    assert response.json()["content"]["structured_content"] is None
    page = response.json()["content"]["pages"][0]
    assert page["ocr_confidence"] is None
    assert page["ocr_line_confidences"] == []


def test_readable_text_absent_for_empty_text_result(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, _ = ocr_fakes
    adapter.outputs = [""]
    upload, _ = _upload_and_audit(client, "empty.png", _image_bytes("PNG"), "image/png")

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    content = response.json()["content"]
    assert content["text"] == ""
    assert content["readable_text"] is None
    assert content["reading_text"] is None


def _pdf_contact_blocks_and_table_bytes() -> bytes:
    """Synthetic single-page offer: two multi-line contact blocks and a table. No real data.

    The left block (x=40) and right block (x=320) each carry several lines so a reading-order
    reconstruction can be checked for "left block fully, then right block fully" rather than only a
    single shared row. Line/company/street names are invented and structurally similar to, but not
    copied from, real documents.
    """
    runs = [
        (40, 770, "Sanierungsbau Perchtoldsdorf GmbH"),
        (320, 770, "Herr Dipl.-Ing. Franz Hubermayr"),
        (40, 752, "Lindenstrasse 42"),
        (320, 752, "Anna Hubermayr geb. Steininger"),
        (40, 734, "2380 Perchtoldsdorf Oesterreich"),
        (320, 734, "Rosengasse 7/12"),
        (40, 716, "Tel: +43 660 1234567"),
        (320, 716, "2340 Moedling Oesterreich"),
        (40, 698, "office@example-contractor.at"),
        (320, 698, "Tel: +43 699 8765432"),
        (40, 680, "UID: ATU12345678"),
        (320, 680, "franz.hubermayr@example.at"),
        (320, 662, "Geburtsdatum: 14.03.1978"),
        (40, 600, "Pos."), (90, 600, "Leistung"), (330, 600, "Menge"),
        (390, 600, "Einheit"), (450, 600, "Einzelpreis"), (530, 600, "Gesamt"),
        (40, 582, "1"), (90, 582, "Abbrucharbeiten Innenwaende"), (330, 582, "45"),
        (390, 582, "m2"), (450, 582, "38,00"), (530, 582, "1710,00"),
        (40, 564, "2"), (90, 564, "Fassadendaemmung"), (330, 564, "180"),
        (390, 564, "m2"), (450, 564, "92,00"), (530, 564, "16560,00"),
    ]
    writer = PdfWriter()
    page = writer.add_blank_page(width=600, height=800)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})}
    )
    data = "".join(f"BT /F1 10 Tf {x} {y} Td ({text}) Tj ET\n" for x, y, text in runs)
    stream = DecodedStreamObject()
    stream.set_data(data.encode("latin-1"))
    page[NameObject("/Contents")] = writer._add_object(stream)
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def test_pii_input_text_generated_for_pdf_text_layer(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, renderer = ocr_fakes
    fixture = _pdf_contact_blocks_and_table_bytes()
    upload, _ = _upload_and_audit(client, "offer.pdf", fixture, "application/pdf")

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    content = response.json()["content"]
    pii_input_text = content["pii_input_text"]
    assert pii_input_text is not None
    assert "[BLOCK: left]" in pii_input_text
    assert "[BLOCK: right]" in pii_input_text
    assert "[TABLE]" in pii_input_text
    assert adapter.calls == []
    assert renderer.calls == []


def test_pii_input_text_left_block_fully_before_right_block(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, _ = _upload_and_audit(
        client, "offer.pdf", _pdf_contact_blocks_and_table_bytes(), "application/pdf"
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    pii_input_text = response.json()["content"]["pii_input_text"]
    left_index = pii_input_text.index("[BLOCK: left]")
    right_index = pii_input_text.index("[BLOCK: right]")
    table_index = pii_input_text.index("[TABLE]")
    # Whole left block precedes the whole right block, which precedes the table — not an
    # X/Y-interleaved dump of both columns.
    assert left_index < right_index < table_index


def test_pii_input_text_keeps_contractor_lines_in_left_block_only(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, _ = _upload_and_audit(
        client, "offer.pdf", _pdf_contact_blocks_and_table_bytes(), "application/pdf"
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    pii_input_text = response.json()["content"]["pii_input_text"]
    left_index = pii_input_text.index("[BLOCK: left]")
    right_index = pii_input_text.index("[BLOCK: right]")
    left_segment = pii_input_text[left_index:right_index]
    right_segment = pii_input_text[right_index:]
    # A left-column line lands in the left block and never among the right block's lines.
    assert "Lindenstrasse 42" in left_segment
    assert "Lindenstrasse 42" not in right_segment


def test_pii_input_text_keeps_customer_lines_in_right_block(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, _ = _upload_and_audit(
        client, "offer.pdf", _pdf_contact_blocks_and_table_bytes(), "application/pdf"
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    pii_input_text = response.json()["content"]["pii_input_text"]
    right_index = pii_input_text.index("[BLOCK: right]")
    table_index = pii_input_text.index("[TABLE]")
    right_segment = pii_input_text[right_index:table_index]
    assert "Herr Dipl.-Ing. Franz Hubermayr" in right_segment
    assert "Anna Hubermayr geb. Steininger" in right_segment


def test_pii_input_text_table_rows_reconstructed_row_wise(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, _ = _upload_and_audit(
        client, "offer.pdf", _pdf_contact_blocks_and_table_bytes(), "application/pdf"
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    pii_input_text = response.json()["content"]["pii_input_text"]
    assert _line_with(pii_input_text, "Pos.", "Leistung", "Gesamt") is not None
    assert _line_with(pii_input_text, "Abbrucharbeiten Innenwaende", "1710,00") is not None


def test_pii_input_text_does_not_change_canonical_text(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    fixture = _pdf_contact_blocks_and_table_bytes()
    upload, _ = _upload_and_audit(client, "offer.pdf", fixture, "application/pdf")

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    content = response.json()["content"]
    expected_canonical = "\n\n".join(
        page.extract_text() or "" for page in PdfReader(BytesIO(fixture)).pages
    )
    assert content["text"] == expected_canonical


def test_pii_input_text_does_not_affect_layout_text_result(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, _ = _upload_and_audit(
        client, "offer.pdf", _pdf_two_column_table_bytes(), "application/pdf"
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    content = response.json()["content"]
    # Both additive fields are produced independently for the same document; pii_input_text does
    # not replace or alter the existing layout_text_result behaviour.
    assert content["pii_input_text"] is not None
    layout = content["layout_text_result"]
    assert layout is not None
    assert _line_with(layout, "AUFTRAGNEHMER", "AUFTRAGGEBER") is not None
    assert _line_with(layout, "Pos.", "Leistung", "Gesamt") is not None


def test_pii_input_text_absent_for_docx(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, _ = _upload_and_audit(
        client,
        "document.docx",
        _docx_bytes("First", "Second"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    assert response.json()["content"]["pii_input_text"] is None


def test_pii_input_text_absent_for_image(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, _ = ocr_fakes
    adapter.outputs = ["Image text"]
    upload, _ = _upload_and_audit(client, "image.png", _image_bytes("PNG"), "image/png")

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    assert response.json()["content"]["pii_input_text"] is None


def test_pii_input_text_marks_ocr_pages_and_page_boundaries(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, _ = ocr_fakes
    adapter.outputs = ["Scanned page two text"]
    upload, _ = _upload_and_audit(
        client, "mixed.pdf", _pdf_pages_bytes(_GOOD_TEXT, None), "application/pdf"
    )

    response = client.post(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 201
    pii_input_text = response.json()["content"]["pii_input_text"]
    assert pii_input_text is not None
    assert "[PAGE 2]" in pii_input_text
    assert "[page 2: pii_input_text not reconstructed]" in pii_input_text
    assert "Scanned page two text" in pii_input_text


def test_legacy_text_artifact_without_pii_input_text_field_remains_valid(
    client: TestClient,
    document_data_dir: Path,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, _ = _upload_and_audit(
        client, "text.pdf", _pdf_pages_bytes("Digital text"), "application/pdf"
    )
    created = client.post(f"/api/documents/{upload['id']}/ocr").json()
    path = _artifact_path(document_data_dir, upload["id"], created["id"])
    payload = json.loads(path.read_text(encoding="utf-8"))
    # Simulate a legacy artifact written before this field existed.
    payload["content"].pop("pii_input_text", None)
    path.write_text(json.dumps(payload), encoding="utf-8")

    response = client.get(f"/api/documents/{upload['id']}/ocr")

    assert response.status_code == 200
    assert response.json()["content"]["pii_input_text"] is None


def test_pdf_text_layer_emits_span_geometry(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, _ = _upload_and_audit(
        client, "text.pdf", _pdf_pages_bytes("Digital text"), "application/pdf"
    )

    content = client.post(f"/api/documents/{upload['id']}/ocr").json()["content"]

    # Technical raw text and its counts remain byte-stable with geometry added.
    assert content["text"] == "Digital text"
    assert content["text_char_count"] == len("Digital text")
    assert content["pages"][0]["text"] == "Digital text"
    assert content["text_geometry_version"] == "1"
    geometry = content["text_geometry"]
    assert geometry is not None
    page = geometry["pages"][0]
    assert page["coordinate_unit"] == "pdf_points"
    assert page["source"] == "pdf_text_layer"
    assert page["status"] == "complete"
    line = page["lines"][0]
    assert content["text"][line["canonical_start"] : line["canonical_end"]] == "Digital text"
    assert line["confidence"] is None
    assert "text_layer_geometry" in geometry["flags"]


def test_mixed_pdf_emits_combined_partial_geometry(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, _ = ocr_fakes
    adapter.outputs = ["Scan two", "Scan three"]
    # Page 2 has OCR polygons; page 3 has none, so its geometry degrades to unsupported.
    adapter.layout_lines = [
        (
            OcrLayoutLine(
                text="Scan two",
                polygon=((10.0, 20.0), (190.0, 20.0), (190.0, 50.0), (10.0, 50.0)),
                confidence=0.88,
            ),
        ),
        (),
    ]
    adapter.image_sizes = [(200, 100), (200, 100)]
    upload, _ = _upload_and_audit(
        client, "mixed.pdf", _pdf_pages_bytes("Digital one", None, None), "application/pdf"
    )

    content = client.post(f"/api/documents/{upload['id']}/ocr").json()["content"]

    assert content["text"] == "Digital one\n\nScan two\n\nScan three"
    geometry = content["text_geometry"]
    assert [page["page_number"] for page in geometry["pages"]] == [1, 2, 3]
    statuses = {page["page_number"]: page["status"] for page in geometry["pages"]}
    assert statuses == {1: "complete", 2: "complete", 3: "unsupported"}
    assert geometry["coverage"] == pytest.approx(2 / 3)
    assert "mixed_geometry" in geometry["flags"]
    assert "partial_geometry" in geometry["flags"]
    # The OCR page's line maps back into the correct canonical region.
    page_two_line = next(
        page for page in geometry["pages"] if page["page_number"] == 2
    )["lines"][0]
    start, end = page_two_line["canonical_start"], page_two_line["canonical_end"]
    assert content["text"][start:end] == "Scan two"
    assert page_two_line["confidence"] == 0.88


def test_image_emits_ocr_span_geometry(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    adapter, _ = ocr_fakes
    adapter.outputs = ["Image text"]
    adapter.layout_lines = [
        (
            OcrLayoutLine(
                text="Image text",
                polygon=((10.0, 20.0), (190.0, 20.0), (190.0, 50.0), (10.0, 50.0)),
                confidence=0.9,
            ),
        )
    ]
    adapter.image_sizes = [(200, 100)]
    upload, _ = _upload_and_audit(client, "image.png", _image_bytes("PNG"), "image/png")

    content = client.post(f"/api/documents/{upload['id']}/ocr").json()["content"]

    assert content["text"] == "Image text"
    geometry = content["text_geometry"]
    assert geometry["pages"][0]["coordinate_unit"] == "image_pixels"
    assert geometry["pages"][0]["status"] == "complete"
    assert geometry["pages"][0]["lines"][0]["confidence"] == 0.9
    assert "ocr_geometry" in geometry["flags"]


def test_docx_has_no_span_geometry(
    client: TestClient,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, _ = _upload_and_audit(
        client,
        "document.docx",
        _docx_bytes("First", "Second"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    content = client.post(f"/api/documents/{upload['id']}/ocr").json()["content"]

    assert content["text_geometry_version"] is None
    assert content["text_geometry"] is None


def test_delete_removes_audit_text_and_quality_artifacts(
    client: TestClient,
    document_data_dir: Path,
    ocr_fakes: tuple[FakeOcrAdapter, FakePdfRenderer],
) -> None:
    upload, _ = _upload_and_audit(
        client, "text.pdf", _pdf_pages_bytes("Text"), "application/pdf"
    )
    assert client.post(f"/api/documents/{upload['id']}/ocr").status_code == 201
    artifact_directory = document_data_dir / str(upload["id"]) / "artifacts"
    assert len(list(artifact_directory.glob("*.json"))) == 3

    response = client.delete(f"/api/documents/{upload['id']}")

    assert response.status_code == 204
    assert not artifact_directory.exists()
