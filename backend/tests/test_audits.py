"""Integration tests for synchronous Audit v1 analysis and artifact persistence."""

from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from PIL import Image
from pydantic import ValidationError
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from app.config import Settings
from app.schemas import AuditContent


@pytest.fixture(autouse=True)
def _allow_larger_audit_fixtures(settings: Settings) -> None:
    settings.max_upload_bytes = 2 * 1024 * 1024


def _upload(client: TestClient, name: str, content: bytes, content_type: str) -> dict[str, object]:
    response = client.post("/api/uploads", files={"file": (name, content, content_type)})
    assert response.status_code == 201
    return response.json()


def _pdf_bytes(*, text: str | None = None) -> bytes:
    return _pdf_pages_bytes(text)


def _pdf_pages_bytes(*page_texts: str | None) -> bytes:
    writer = PdfWriter()
    for text in page_texts:
        page = writer.add_blank_page(width=200, height=200)
        if text is not None:
            font = DictionaryObject(
                {
                    NameObject("/Type"): NameObject("/Font"),
                    NameObject("/Subtype"): NameObject("/Type1"),
                    NameObject("/BaseFont"): NameObject("/Helvetica"),
                }
            )
            page[NameObject("/Resources")] = DictionaryObject(
                {
                    NameObject("/Font"): DictionaryObject(
                        {NameObject("/F1"): writer._add_object(font)}
                    )
                }
            )
            stream = DecodedStreamObject()
            stream.set_data(f"BT /F1 12 Tf 10 100 Td ({text}) Tj ET".encode())
            page[NameObject("/Contents")] = writer._add_object(stream)
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def _docx_bytes(*paragraphs: str) -> bytes:
    document = DocxDocument()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _docx_with_table_bytes() -> bytes:
    """A paragraph, a 2x2 table, then a paragraph — to exercise ordered table extraction."""
    document = DocxDocument()
    document.add_paragraph("Intro")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "R1C1"
    table.rows[0].cells[1].text = "R1C2"
    table.rows[1].cells[0].text = "R2C1"
    table.rows[1].cells[1].text = "R2C2"
    document.add_paragraph("Outro")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _image_bytes(image_format: str, size: tuple[int, int]) -> bytes:
    buffer = BytesIO()
    Image.new("RGB", size, color="white").save(buffer, format=image_format)
    return buffer.getvalue()


def _metadata(document_data_dir: Path, document_id: str) -> tuple[Path, dict[str, object]]:
    path = document_data_dir / document_id / "document.json"
    return path, json.loads(path.read_text(encoding="utf-8"))


def _artifact_directory(document_data_dir: Path, document_id: object) -> Path:
    return document_data_dir / str(document_id) / "artifacts"


def test_audits_pdf_with_text_layer(
    client: TestClient, document_data_dir: Path
) -> None:
    upload = _upload(client, "text.pdf", _pdf_bytes(text="Audit text"), "application/pdf")

    response = client.post(f"/api/documents/{upload['id']}/audit")

    assert response.status_code == 201
    artifact = response.json()
    content = artifact["content"]
    assert artifact["artifact_type"] == "audit_result"
    assert artifact["station"] == "audit"
    assert artifact["media_type"] == "application/json"
    assert artifact["input_artifact_id"] == upload["original_artifact"]["id"]
    assert content["document_id"] == upload["id"]
    assert content["input_artifact_id"] == upload["original_artifact"]["id"]
    assert content["detected_mime_type"] == "application/pdf"
    assert content["audit_version"] == "1"
    assert content["document_kind"] == "pdf"
    assert content["page_count"] == 1
    assert content["has_text_layer"] is True
    assert content["text_char_count"] == len("Audit text")
    page = content["pages"][0]
    assert page["page_number"] == 1
    assert page["text_char_count"] == len("Audit text")
    assert page["has_text_layer"] is True
    # "Audit text" is plausible but too short to be confident, so it stays on the text layer.
    assert page["needs_ocr"] is False
    assert page["recommended_text_source"] == "text_layer"
    assert page["text_quality_status"] in {"GOOD_TEXT_LAYER", "LOW_CONFIDENCE_TEXT_LAYER"}
    assert isinstance(page["text_quality_score"], int)
    assert isinstance(page["text_quality_reasons"], list)
    assert content["flags"] == ["pdf_has_text_layer"]
    assert content["tool_versions"]["pypdf"]
    artifact_path = _artifact_directory(document_data_dir, upload["id"]) / f"{artifact['id']}.json"
    assert artifact_path.is_file()


def test_audits_pdf_without_text_layer(client: TestClient) -> None:
    upload = _upload(client, "blank.pdf", _pdf_bytes(), "application/pdf")

    response = client.post(f"/api/documents/{upload['id']}/audit")

    assert response.status_code == 201
    content = response.json()["content"]
    assert content["page_count"] == 1
    assert content["has_text_layer"] is False
    assert content["text_char_count"] == 0
    page = content["pages"][0]
    assert page["has_text_layer"] is False
    assert page["text_quality_status"] == "EMPTY_TEXT_LAYER"
    assert page["needs_ocr"] is True
    assert page["recommended_text_source"] == "ocr"
    # An empty page still routes to OCR, so the additive summary flag is present.
    assert content["flags"] == ["pdf_no_text_layer", "pdf_pages_need_ocr"]


def test_audits_mixed_pdf_text_layers_per_page(client: TestClient) -> None:
    page_text = "Digital first page"
    upload = _upload(
        client,
        "mixed.pdf",
        _pdf_pages_bytes(page_text, None, None),
        "application/pdf",
    )

    response = client.post(f"/api/documents/{upload['id']}/audit")

    assert response.status_code == 201
    content = response.json()["content"]
    assert content["page_count"] == 3
    assert content["has_text_layer"] is True
    assert content["text_char_count"] == sum(
        page["text_char_count"] for page in content["pages"]
    )
    pages = content["pages"]
    assert [page["has_text_layer"] for page in pages] == [True, False, False]
    assert [page["text_char_count"] for page in pages] == [len(page_text), 0, 0]
    # The text page keeps its layer; the two empty pages are routed to OCR.
    assert [page["needs_ocr"] for page in pages] == [False, True, True]
    assert [page["text_quality_status"] for page in pages] == [
        "LOW_CONFIDENCE_TEXT_LAYER",
        "EMPTY_TEXT_LAYER",
        "EMPTY_TEXT_LAYER",
    ]
    assert "pdf_pages_need_ocr" in content["flags"]


def test_audit_good_page_does_not_need_ocr(client: TestClient) -> None:
    good_text = "The quick brown fox jumps over the lazy dog near the calm winding river today"
    upload = _upload(client, "good.pdf", _pdf_bytes(text=good_text), "application/pdf")

    response = client.post(f"/api/documents/{upload['id']}/audit")

    assert response.status_code == 201
    content = response.json()["content"]
    page = content["pages"][0]
    assert page["text_quality_status"] == "GOOD_TEXT_LAYER"
    assert page["needs_ocr"] is False
    assert page["recommended_text_source"] == "text_layer"
    assert page["text_quality_score"] >= 80
    assert content["flags"] == ["pdf_has_text_layer"]


def test_audit_broken_text_layer_page_needs_ocr(client: TestClient) -> None:
    # Many characters, no letters, interior symbols: the broken/encoded text-layer signature.
    garbage = "1#2 3%4 5@6 7|8 9^0 2&3 4*5 6~7 8<9 0>1 2?3 4=5 6#7 8%9 0@1 2|3 4^5 6&7 8<1 0>2"
    upload = _upload(client, "broken.pdf", _pdf_bytes(text=garbage), "application/pdf")

    response = client.post(f"/api/documents/{upload['id']}/audit")

    assert response.status_code == 201
    content = response.json()["content"]
    page = content["pages"][0]
    # The page formally has a text layer, but its content is unusable and must be OCR'd.
    assert page["has_text_layer"] is True
    assert page["text_quality_status"] == "BROKEN_TEXT_LAYER"
    assert page["needs_ocr"] is True
    assert page["recommended_text_source"] == "ocr"
    assert "pdf_broken_text_layer" in content["flags"]
    assert "pdf_pages_need_ocr" in content["flags"]


@pytest.mark.parametrize(
    ("field", "value"),
    [
        ("page_count", 1),
        (
            "pages",
            [
                {"page_number": 1, "text_char_count": 4, "has_text_layer": True},
                {"page_number": 3, "text_char_count": 0, "has_text_layer": False},
            ],
        ),
        ("text_char_count", 99),
        ("has_text_layer", False),
    ],
)
def test_pdf_audit_content_rejects_inconsistent_page_summary(
    field: str, value: object
) -> None:
    content: dict[str, object] = {
        "document_id": "a" * 32,
        "input_artifact_id": "b" * 32,
        "detected_mime_type": "application/pdf",
        "document_kind": "pdf",
        "page_count": 2,
        "has_text_layer": True,
        "text_char_count": 4,
        "pages": [
            {"page_number": 1, "text_char_count": 4, "has_text_layer": True},
            {"page_number": 2, "text_char_count": 0, "has_text_layer": False},
        ],
    }
    content[field] = value

    with pytest.raises(ValidationError):
        AuditContent.model_validate(content)


def test_audits_docx_body_paragraphs(client: TestClient) -> None:
    upload = _upload(
        client,
        "report.docx",
        _docx_bytes("First", "Second paragraph"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    response = client.post(f"/api/documents/{upload['id']}/audit")

    assert response.status_code == 201
    content = response.json()["content"]
    assert content["document_kind"] == "docx"
    assert content["paragraph_count"] == 2
    # Shared extraction joins body paragraphs with a newline (see docx_extraction).
    assert content["text_char_count"] == len("First\nSecond paragraph")
    assert content["has_text_layer"] is True
    assert content["tool_versions"]["python-docx"]


def test_audits_docx_counts_table_content(client: TestClient) -> None:
    upload = _upload(
        client,
        "tables.docx",
        _docx_with_table_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    response = client.post(f"/api/documents/{upload['id']}/audit")

    assert response.status_code == 201
    content = response.json()["content"]
    assert content["document_kind"] == "docx"
    # Table cell text ("Intro\nR1C1\tR1C2\nR2C1\tR2C2\nOutro") must be counted, not just the
    # body paragraphs — paragraph-only extraction would report far fewer characters.
    expected = "Intro\nR1C1\tR1C2\nR2C1\tR2C2\nOutro"
    assert content["text_char_count"] == len(expected)
    assert content["text_char_count"] > len("IntroOutro")


@pytest.mark.parametrize(
    ("extension", "mime_type", "image_format"),
    [("png", "image/png", "PNG"), ("jpg", "image/jpeg", "JPEG")],
)
def test_audits_image_dimensions(
    client: TestClient, extension: str, mime_type: str, image_format: str
) -> None:
    upload = _upload(client, f"image.{extension}", _image_bytes(image_format, (13, 17)), mime_type)

    response = client.post(f"/api/documents/{upload['id']}/audit")

    assert response.status_code == 201
    content = response.json()["content"]
    assert content["document_kind"] == "image"
    assert content["image_format"] == image_format
    assert (content["width"], content["height"]) == (13, 17)
    assert content["has_text_layer"] is False
    assert content["tool_versions"]["Pillow"]


def test_dispatches_by_artifact_mime_not_extension(
    client: TestClient, upload_dir: Path, document_data_dir: Path
) -> None:
    upload = _upload(client, "image.png", _image_bytes("PNG", (7, 9)), "image/png")
    document_id = str(upload["id"])
    metadata_path, metadata = _metadata(document_data_dir, document_id)
    original = metadata["original_artifact"]
    assert isinstance(original, dict)
    old_path = upload_dir / str(original["storage_filename"])
    new_path = upload_dir / f"{document_id}.pdf"
    old_path.rename(new_path)
    metadata["extension"] = "pdf"
    original["storage_filename"] = new_path.name
    metadata_path.write_text(json.dumps(metadata), encoding="utf-8")

    response = client.post(f"/api/documents/{document_id}/audit")

    assert response.status_code == 201
    assert response.json()["content"]["document_kind"] == "image"


def test_legacy_metadata_without_original_artifact_returns_409(
    client: TestClient, document_data_dir: Path
) -> None:
    upload = _upload(client, "blank.pdf", _pdf_bytes(), "application/pdf")
    document_id = str(upload["id"])
    metadata_path, metadata = _metadata(document_data_dir, document_id)
    metadata.pop("original_artifact")
    metadata.pop("sha256")
    metadata.pop("detected_mime_type")
    metadata_path.write_text(json.dumps(metadata), encoding="utf-8")

    response = client.post(f"/api/documents/{document_id}/audit")

    assert response.status_code == 409
    assert list(_artifact_directory(document_data_dir, document_id).iterdir()) == []


def test_hash_mismatch_returns_409_without_artifact(
    client: TestClient, upload_dir: Path, document_data_dir: Path
) -> None:
    upload = _upload(client, "blank.pdf", _pdf_bytes(), "application/pdf")
    original = upload["original_artifact"]
    assert isinstance(original, dict)
    (upload_dir / str(original["storage_filename"])).write_bytes(b"tampered")

    response = client.post(f"/api/documents/{upload['id']}/audit")

    assert response.status_code == 409
    assert list(_artifact_directory(document_data_dir, upload["id"]).iterdir()) == []


def test_missing_original_file_returns_409_without_artifact(
    client: TestClient, upload_dir: Path, document_data_dir: Path
) -> None:
    upload = _upload(client, "blank.pdf", _pdf_bytes(), "application/pdf")
    original = upload["original_artifact"]
    assert isinstance(original, dict)
    (upload_dir / str(original["storage_filename"])).unlink()

    response = client.post(f"/api/documents/{upload['id']}/audit")

    assert response.status_code == 409
    assert list(_artifact_directory(document_data_dir, upload["id"]).iterdir()) == []


def test_corrupt_file_returns_422(client: TestClient) -> None:
    upload = _upload(client, "broken.pdf", b"%PDF-1.4 not actually a PDF", "application/pdf")

    response = client.post(f"/api/documents/{upload['id']}/audit")

    assert response.status_code == 422


def test_unsupported_artifact_mime_returns_422(
    client: TestClient, document_data_dir: Path
) -> None:
    upload = _upload(client, "blank.pdf", _pdf_bytes(), "application/pdf")
    document_id = str(upload["id"])
    metadata_path, metadata = _metadata(document_data_dir, document_id)
    original = metadata["original_artifact"]
    assert isinstance(original, dict)
    original["mime_type"] = "application/octet-stream"
    metadata["detected_mime_type"] = "application/octet-stream"
    metadata["content_type"] = "application/octet-stream"
    metadata_path.write_text(json.dumps(metadata), encoding="utf-8")

    response = client.post(f"/api/documents/{document_id}/audit")

    assert response.status_code == 422


def test_get_without_audit_returns_404(client: TestClient) -> None:
    upload = _upload(client, "blank.pdf", _pdf_bytes(), "application/pdf")

    response = client.get(f"/api/documents/{upload['id']}/audit")

    assert response.status_code == 404


def test_unknown_document_returns_404(client: TestClient) -> None:
    document_id = "a" * 32

    assert client.post(f"/api/documents/{document_id}/audit").status_code == 404
    assert client.get(f"/api/documents/{document_id}/audit").status_code == 404


def test_get_returns_latest_audit(client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    upload = _upload(client, "blank.pdf", _pdf_bytes(), "application/pdf")
    timestamps = iter(["2026-07-01T10:00:00.000001Z", "2026-07-01T10:00:00.000002Z"])
    monkeypatch.setattr("app.services.audit_service._now_utc_iso", lambda: next(timestamps))
    first = client.post(f"/api/documents/{upload['id']}/audit")
    second = client.post(f"/api/documents/{upload['id']}/audit")

    response = client.get(f"/api/documents/{upload['id']}/audit")

    assert first.status_code == 201
    assert second.status_code == 201
    assert first.json()["id"] != second.json()["id"]
    assert response.status_code == 200
    assert response.json()["id"] == second.json()["id"]


def test_artifact_finalize_failure_removes_partial_file(
    client: TestClient, document_data_dir: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    upload = _upload(client, "blank.pdf", _pdf_bytes(), "application/pdf")
    original_replace = Path.replace

    def fail_artifact_replace(path: Path, target: Path) -> Path:
        if path.name.endswith(".json.part") and "artifacts" in path.parts:
            raise OSError("simulated artifact finalize failure")
        return original_replace(path, target)

    monkeypatch.setattr(Path, "replace", fail_artifact_replace)

    with pytest.raises(OSError, match="simulated artifact finalize failure"):
        client.post(f"/api/documents/{upload['id']}/audit")

    artifact_directory = _artifact_directory(document_data_dir, upload["id"])
    assert artifact_directory.is_dir()
    assert list(artifact_directory.iterdir()) == []


def test_delete_removes_audit_artifact_directory(
    client: TestClient, upload_dir: Path, document_data_dir: Path
) -> None:
    upload = _upload(client, "blank.pdf", _pdf_bytes(), "application/pdf")
    assert client.post(f"/api/documents/{upload['id']}/audit").status_code == 201
    artifact_directory = _artifact_directory(document_data_dir, upload["id"])
    assert artifact_directory.is_dir()

    response = client.delete(f"/api/documents/{upload['id']}")

    assert response.status_code == 204
    assert not artifact_directory.exists()
    assert list(upload_dir.iterdir()) == []
