"""Shared, deterministic DOCX text extraction used by Audit and OCR.

``python-docx`` exposes body paragraphs via ``document.paragraphs`` but not the paragraphs
nested inside table cells, so paragraph-only extraction silently drops table content — which is
frequently the bulk of a document. This helper walks the document body in document order,
captures both paragraphs and tables, and adds section header/footer text. It returns one
deterministic string so the Audit and OCR/Text workstations share a single DOCX interpretation
and cannot diverge.

Scope (v1, ``python-docx`` only — no new tools, no OCR, no layout reconstruction):

- body paragraphs, in document order;
- tables: rows joined by newline, cells within a row joined by a tab; horizontally merged cells
  (which ``python-docx`` yields repeatedly) are emitted once;
- section headers/footers that are actually defined (not merely linked to a previous section);
- textboxes/shapes and nested tables are out of scope — not present in the target corpus and not
  cleanly reachable via ``python-docx``.

The output is a pure function of the input bytes, so downstream PII offsets stay stable.
"""

from __future__ import annotations

import re
from collections.abc import Iterator

from docx.document import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

# Collapse runs of three or more newlines to a single blank line so documents with many empty
# paragraphs stay readable without dropping meaningful paragraph spacing.
_BLANK_LINE_RUN = re.compile(r"\n{3,}")


def extract_docx_text(document: Document) -> str:
    """Return the document's text with tables and headers/footers, in document order."""
    lines: list[str] = []
    lines.extend(_section_lines(document, "header"))
    lines.extend(_body_lines(document))
    lines.extend(_section_lines(document, "footer"))
    return _BLANK_LINE_RUN.sub("\n\n", "\n".join(lines)).strip()


def _body_lines(document: Document) -> list[str]:
    lines: list[str] = []
    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            lines.append(block.text)
        else:
            lines.extend(_table_lines(block))
    return lines


def _iter_block_items(document: Document) -> Iterator[Paragraph | Table]:
    """Yield body-level paragraphs and tables in document order.

    ``document.paragraphs`` and ``document.tables`` each flatten one kind and lose the
    interleaving; iterating the body's XML children preserves the original order.
    """
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _table_lines(table: Table) -> list[str]:
    lines: list[str] = []
    for row in table.rows:
        seen: set[int] = set()
        cells: list[str] = []
        for cell in row.cells:
            # Horizontally merged cells repeat the same underlying <w:tc>; emit it once.
            key = id(cell._tc)
            if key in seen:
                continue
            seen.add(key)
            cells.append(cell.text)
        lines.append("\t".join(cells))
    return lines


def _section_lines(document: Document, part: str) -> list[str]:
    lines: list[str] = []
    for section in document.sections:
        header_footer = getattr(section, part)
        # A linked header/footer inherits from an earlier section; skip it to avoid duplication.
        if header_footer.is_linked_to_previous:
            continue
        text = "\n".join(paragraph.text for paragraph in header_footer.paragraphs)
        if text.strip():
            lines.append(text)
    return lines
